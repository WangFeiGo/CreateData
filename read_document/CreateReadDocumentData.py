#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Author   : wangfei
# @Time     : 2020/12/8 9:54
# @File     : CreateReadDocumentData.py
# @Software : PyCharm

# 添加环境变量
import os,sys
ENV_DIR = os.path.abspath(os.path.join(os.getcwd(), ".."))
sys.path.append(ENV_DIR)

import docx
import os
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
import random
from common import ReadJSON

class CreateReadDocumentData():

    # 如果需要手输入关键词，则放开
    # key_words = input("请输入对应的关键字(捕获关键字):")

    def __init__(self):

        # self.file = docx.Document(r"C:\Users\Administrator\AppData\Local\Programs\Python\Python36\Lib\site-packages\docx\templates\default.docx")
        self.file = docx.Document()
        self.data = ReadJSON.readJSON(os.path.abspath('..') + "\\read_document\\config" + "\\data.json")
        self.big_title = random.choice(self.data["big_title"])
        self.small_title = random.choice(self.data["small_title"])
        self.article_title = random.choice(self.data["article_title"])
        self.paragraph1 = random.choice(self.data["paragraph1"])
        self.paragraph2 = random.choice(self.data["paragraph2"])
        self.paragraph3 = random.choice(self.data["paragraph3"])
        # 随机生成关键词
        self.key_words = random.choice(self.data["key_word"])
        # 秘密关键字读取
        self.secret_id = random.choice(self.data["secret"])
        # 对齐方式
        self.alignment = random.choice([WD_PARAGRAPH_ALIGNMENT.CENTER,
                                         WD_PARAGRAPH_ALIGNMENT.LEFT,
                                         WD_PARAGRAPH_ALIGNMENT.RIGHT])
        # 页眉内容
        self.head_content = random.choice(["页眉",self.secret_id])

    def create_header(self):
        # 页眉
        self.header = self.file.sections[0].header
        p = self.header.paragraphs[0]
        p.alignment = self.alignment
        run = p.add_run(self.head_content)
        run.font.size = Pt(14)

    def create_secret(self):
        # 秘密关键字
        p = self.file.add_paragraph()
        p.alignment = self.alignment
        run = p.add_run(self.secret_id)
        run.font.name = u'华文中宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文中宋')
        run.font.size = Pt(18)

    def create_big_title(self):
        # 大标题
        p = self.file.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(self.big_title)
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.font.name=u'华文中宋'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'华文中宋')
        run.font.size = Pt(48)
        if self.head_content == "页眉":
            self.create_secret()
        # 空三行
        self.file.add_paragraph("\n")
        self.file.add_paragraph("\n")
        self.file.add_paragraph("\n")

    def create_small_title(self):
        # 小标题
        p = self.file.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(self.small_title)
        run.font.name=u'仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
        run.font.size = Pt(18)
        # 分割线
        self.file.add_paragraph("_________________________________________________________________________________________________________")

    def create_document_title(self):
        # 文章标题
        p = self.file.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.article_title = self.article_title.replace("Exchange",self.key_words)
        run = p.add_run(self.article_title)
        run.font.name=u'仿宋_GB2312'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
        run.font.size = Pt(20)
        run.font.bold = True

    def create_document_content(self):
        # 文章内容
        datas = [self.paragraph1,self.paragraph2,self.paragraph3]
        for data in datas:
            p = self.file.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.first_line_indent = Inches(0.5)
            run = p.add_run(data.replace("Exchange",self.key_words))
            run.font.name=u'仿宋_GB2312'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
            run.font.size = Pt(16)

    def save(self,documentpath):
        self.file.save(os.path.abspath('..') + documentpath + ".doc")

if __name__ == "__main__":

    times = int(input("请输入要造的文件数量(阿拉伯数字):"))

    for i in range(times):
        run = CreateReadDocumentData()
        run.create_header()
        run.create_big_title()
        run.create_small_title()
        run.create_document_title()
        run.create_document_content()
        run.save("\\read_document\\datas\\word\\" + str(i) + "_" + run.article_title + "_" + run.key_words)